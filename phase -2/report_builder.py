from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from sklearn.ensemble import RandomForestRegressor
from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.model_selection import LeaveOneGroupOut
from xgboost import XGBRegressor

import research_pipeline
from research_pipeline import (
    EARLY_EXPONENT_SELECTED,
    LATE_EXPONENT_SELECTED,
    OUTPUTS,
    PHASE2,
    PLOTS,
    fit_common_power,
    fit_fixed_power_amplitudes,
    load_nasa_case_level,
    load_nuaa_run_level,
    load_phm_cut_level,
)
from tool_life_simulator import MillingToolLifeModel


REPORT_PATH = OUTPUTS / "phase2_research_report.docx"


def bootstrap_exponents(
    df: pd.DataFrame,
    group_col: str,
    time_col: str,
    wear_col: str,
    n_boot: int,
    seed: int,
    initial_guess_log_a: float,
    initial_guess_m: float,
) -> pd.DataFrame:
    rng = np.random.default_rng(seed)
    groups = list(pd.Series(df[group_col].unique()).astype(str))
    rows = []

    for i in range(n_boot):
        sampled = rng.choice(groups, size=len(groups), replace=True)
        boot_frames = []
        for j, group in enumerate(sampled):
            part = df[df[group_col].astype(str) == group].copy()
            part[group_col] = f"{group}_boot_{j}"
            boot_frames.append(part)

        boot_df = pd.concat(boot_frames, ignore_index=True)
        fit = fit_common_power(
            df=boot_df,
            group_col=group_col,
            time_col=time_col,
            wear_col=wear_col,
            initial_guess_log_a=initial_guess_log_a,
            initial_guess_m=initial_guess_m,
        )
        rows.append({"bootstrap_id": i, "exponent": fit.exponent})

    return pd.DataFrame(rows)


def equation_predictions_for_fold(train_df: pd.DataFrame, test_df: pd.DataFrame) -> np.ndarray:
    rows = []
    for exp, group_df in train_df.groupby("experiment_tag"):
        x = np.power(np.maximum(group_df["relative_time_min"].to_numpy(dtype=float), 1e-9), EARLY_EXPONENT_SELECTED)
        y = group_df["delta_wear_mm"].to_numpy(dtype=float)
        amplitude = float(np.dot(x, y) / np.dot(x, x))
        rows.append(
            {
                "experiment_tag": exp,
                "speed_rpm": float(group_df["speed_rpm"].iloc[0]),
                "feed_mm_tooth": float(group_df["feed_mm_tooth"].iloc[0]),
                "depth_mm": float(group_df["depth_mm"].iloc[0]),
                "amplitude": amplitude,
            }
        )

    amp_df = pd.DataFrame(rows)
    reg = LinearRegression().fit(
        np.log(
            amp_df[["speed_rpm", "feed_mm_tooth", "depth_mm"]].to_numpy(dtype=float)
            / np.array([1800.0, 0.05, 3.0])
        ),
        np.log(amp_df["amplitude"].to_numpy(dtype=float)),
    )

    x_test = np.log(
        test_df[["speed_rpm", "feed_mm_tooth", "depth_mm"]].to_numpy(dtype=float)
        / np.array([1800.0, 0.05, 3.0])
    )
    amplitude_hat = np.exp(reg.predict(x_test))
    return test_df["initial_wear_mm"].to_numpy(dtype=float) + amplitude_hat * np.power(
        np.maximum(test_df["relative_time_min"].to_numpy(dtype=float), 1e-9),
        EARLY_EXPONENT_SELECTED,
    )


def benchmark_nuaa_models(nuaa_runs: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    logo = LeaveOneGroupOut()
    groups = nuaa_runs["experiment_tag"].astype(str).to_numpy()
    y = nuaa_runs["wear_mm"].to_numpy(dtype=float)

    feature_df = nuaa_runs[
        ["relative_time_min", "speed_rpm", "feed_mm_tooth", "depth_mm", "initial_wear_mm"]
    ].copy()
    feature_df["log_time"] = np.log1p(feature_df["relative_time_min"])
    feature_df = feature_df[
        ["relative_time_min", "log_time", "speed_rpm", "feed_mm_tooth", "depth_mm", "initial_wear_mm"]
    ]

    models = {
        "Equation": None,
        "Linear Regression": LinearRegression(),
        "Random Forest": RandomForestRegressor(
            n_estimators=400,
            random_state=42,
            min_samples_leaf=2,
        ),
        "XGBoost": XGBRegressor(
            n_estimators=350,
            max_depth=4,
            learning_rate=0.05,
            subsample=0.9,
            colsample_bytree=0.9,
            random_state=42,
            objective="reg:squarederror",
        ),
    }

    prediction_rows = []
    metric_rows = []

    for model_name, model in models.items():
        pred = np.zeros(len(nuaa_runs), dtype=float)
        for train_idx, test_idx in logo.split(feature_df, y, groups):
            train_df = nuaa_runs.iloc[train_idx].copy()
            test_df = nuaa_runs.iloc[test_idx].copy()
            if model_name == "Equation":
                pred[test_idx] = equation_predictions_for_fold(train_df=train_df, test_df=test_df)
            else:
                model.fit(feature_df.iloc[train_idx], y[train_idx])
                pred[test_idx] = model.predict(feature_df.iloc[test_idx])

        rmse = float(np.sqrt(mean_squared_error(y, pred)))
        mae = float(mean_absolute_error(y, pred))
        r2 = float(r2_score(y, pred))
        metric_rows.append({"model": model_name, "rmse_mm": rmse, "mae_mm": mae, "r2": r2})

        frame = nuaa_runs[["experiment_tag", "relative_time_min", "wear_mm"]].copy()
        frame["model"] = model_name
        frame["predicted_wear_mm"] = pred
        prediction_rows.append(frame)

    metrics_df = pd.DataFrame(metric_rows).sort_values(["rmse_mm", "mae_mm"])
    predictions_df = pd.concat(prediction_rows, ignore_index=True)
    return metrics_df, predictions_df


def plot_bootstrap_results(
    early_nuaa: pd.DataFrame,
    early_phm: pd.DataFrame,
    late_nasa: pd.DataFrame,
    output_path: Path,
) -> None:
    fig, ax = plt.subplots(figsize=(8, 5.5))
    data = [early_nuaa["exponent"], early_phm["exponent"], late_nasa["exponent"]]
    labels = ["NUAA Early", "PHM Early", "NASA Late"]
    parts = ax.violinplot(data, showmeans=True, showextrema=True)
    colors = ["#4E79A7", "#59A14F", "#E15759"]
    for body, color in zip(parts["bodies"], colors):
        body.set_facecolor(color)
        body.set_alpha(0.5)

    ax.set_xticks([1, 2, 3], labels)
    ax.set_ylabel("Exponent")
    ax.set_title("Bootstrap Stability of Fitted Wear Exponents")
    ax.grid(alpha=0.25, axis="y")
    fig.tight_layout()
    fig.savefig(output_path, dpi=180)
    plt.close(fig)


def plot_benchmark_results(metrics_df: pd.DataFrame, output_path: Path) -> None:
    fig, axes = plt.subplots(1, 3, figsize=(14, 4.8))
    ordered = metrics_df.copy()

    axes[0].bar(ordered["model"], ordered["rmse_mm"], color="#4E79A7")
    axes[0].set_title("RMSE")
    axes[0].set_ylabel("mm")
    axes[0].tick_params(axis="x", rotation=25)

    axes[1].bar(ordered["model"], ordered["mae_mm"], color="#59A14F")
    axes[1].set_title("MAE")
    axes[1].set_ylabel("mm")
    axes[1].tick_params(axis="x", rotation=25)

    axes[2].bar(ordered["model"], ordered["r2"], color="#E15759")
    axes[2].set_title("R²")
    axes[2].tick_params(axis="x", rotation=25)

    for ax in axes:
        ax.grid(alpha=0.25, axis="y")

    fig.suptitle("NUAA Leave-One-Experiment-Out Benchmark")
    fig.tight_layout(rect=(0, 0, 1, 0.95))
    fig.savefig(output_path, dpi=180)
    plt.close(fig)


def plot_parity_results(predictions_df: pd.DataFrame, output_path: Path) -> None:
    keep_models = ["Equation", "Linear Regression"]
    fig, axes = plt.subplots(1, 2, figsize=(11, 5), sharex=True, sharey=True)
    for ax, model_name in zip(axes, keep_models):
        part = predictions_df[predictions_df["model"] == model_name]
        ax.scatter(part["wear_mm"], part["predicted_wear_mm"], s=24, alpha=0.8)
        lo = min(float(part["wear_mm"].min()), float(part["predicted_wear_mm"].min()))
        hi = max(float(part["wear_mm"].max()), float(part["predicted_wear_mm"].max()))
        ax.plot([lo, hi], [lo, hi], linestyle="--", color="black", linewidth=1)
        ax.set_title(model_name)
        ax.set_xlabel("Observed wear (mm)")
        ax.set_ylabel("Predicted wear (mm)")
        ax.grid(alpha=0.25)

    fig.suptitle("Parity Check: Explicit Equation vs Best ML Baseline")
    fig.tight_layout(rect=(0, 0, 1, 0.95))
    fig.savefig(output_path, dpi=180)
    plt.close(fig)


def plot_sensitivity_curves(model: MillingToolLifeModel, output_path: Path) -> None:
    fig, axes = plt.subplots(1, 3, figsize=(15, 4.6))
    baseline = {"speed": 1800.0, "feed": 0.05, "depth": 3.0}

    speeds = np.linspace(1700, 1900, 30)
    life = [
        model.life_to_threshold_minutes(
            speed_rpm=float(s),
            feed_mm_tooth=baseline["feed"],
            depth_mm=baseline["depth"],
            threshold_wear_mm=0.30,
            initial_wear_mm=0.05,
            calibration_factor=1.0,
            material_family="generic",
        )
        for s in speeds
    ]
    axes[0].plot(speeds, life, color="#4E79A7")
    axes[0].set_title("Speed Sensitivity")
    axes[0].set_xlabel("Speed (rpm)")
    axes[0].set_ylabel("Predicted life (min)")

    feeds = np.linspace(0.045, 0.055, 30)
    life = [
        model.life_to_threshold_minutes(
            speed_rpm=baseline["speed"],
            feed_mm_tooth=float(f),
            depth_mm=baseline["depth"],
            threshold_wear_mm=0.30,
            initial_wear_mm=0.05,
            calibration_factor=1.0,
            material_family="generic",
        )
        for f in feeds
    ]
    axes[1].plot(feeds, life, color="#59A14F")
    axes[1].set_title("Feed Sensitivity")
    axes[1].set_xlabel("Feed per tooth (mm/tooth)")

    depths = np.linspace(2.5, 3.5, 30)
    life = [
        model.life_to_threshold_minutes(
            speed_rpm=baseline["speed"],
            feed_mm_tooth=baseline["feed"],
            depth_mm=float(d),
            threshold_wear_mm=0.30,
            initial_wear_mm=0.05,
            calibration_factor=1.0,
            material_family="generic",
        )
        for d in depths
    ]
    axes[2].plot(depths, life, color="#E15759")
    axes[2].set_title("Depth Sensitivity")
    axes[2].set_xlabel("Axial depth (mm)")

    for ax in axes:
        ax.grid(alpha=0.25)

    fig.suptitle("Simulator Sensitivity Around the Baseline Condition")
    fig.tight_layout(rect=(0, 0, 1, 0.95))
    fig.savefig(output_path, dpi=180)
    plt.close(fig)


def plot_late_stage_materials(nasa_ratio_df: pd.DataFrame, output_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7, 5))
    groups = ["cast_iron", "steel"]
    data = [nasa_ratio_df[nasa_ratio_df["material_family"] == g]["late_to_early_ratio"] for g in groups]
    ax.boxplot(data, tick_labels=groups)
    for i, g in enumerate(groups, start=1):
        y = nasa_ratio_df[nasa_ratio_df["material_family"] == g]["late_to_early_ratio"].to_numpy()
        x = np.full_like(y, i, dtype=float) + np.linspace(-0.08, 0.08, len(y))
        ax.scatter(x, y, alpha=0.7, s=30)
    ax.set_ylabel("Late/early amplitude ratio")
    ax.set_title("NASA Late-Stage Acceleration by Material Family")
    ax.grid(alpha=0.25, axis="y")
    fig.tight_layout()
    fig.savefig(output_path, dpi=180)
    plt.close(fig)


def aggregate_dataset_summary(dataset_profiles: pd.DataFrame) -> pd.DataFrame:
    return (
        dataset_profiles.groupby("dataset", as_index=False)
        .agg(
            series_count=("series_id", "count"),
            total_points=("points", "sum"),
            min_wear_mm=("start_wear_mm", "min"),
            max_wear_mm=("end_wear_mm", "max"),
            max_time=("end_time", "max"),
        )
        .sort_values("dataset")
    )


def add_df_table(document: Document, df: pd.DataFrame, decimals: int | None = None) -> None:
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float) and decimals is not None:
                cells[i].text = f"{value:.{decimals}f}"
            else:
                cells[i].text = str(value)


def add_figure(document: Document, image_path: Path, caption: str, width_in: float = 6.5) -> None:
    document.add_picture(str(image_path), width=Inches(width_in))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True


def build_report(
    model_summary: dict[str, object],
    dataset_summary: pd.DataFrame,
    key_results: pd.DataFrame,
    benchmark_metrics: pd.DataFrame,
    late_ratio_summary: pd.DataFrame,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Phase-2 Research Report\nData-Backed Milling Tool-Life Equation and Simulation Engine")
    run.bold = True
    run.font.size = Pt(18)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Prepared on {date.today().isoformat()}").italic = True

    document.add_paragraph(
        "This report summarizes the phase-2 research workflow built from the NUAA orthogonal milling bundle, "
        "the PHM 2010 milling challenge dataset, and the NASA milling dataset. The objective was to derive "
        "a data-backed analytical wear equation that can be inverted into a practical tool-life simulator."
    )

    document.add_heading("1. Executive Summary", level=1)
    summary_points = [
        f"The fitted early-stage wear exponent is {model_summary['evidence']['nuaa_early_exponent_raw']:.4f} on NUAA and {model_summary['evidence']['phm_early_exponent_raw']:.4f} on PHM 2010, which justifies the selected operating value of 0.64.",
        f"The fitted late-stage wear exponent on the NASA run-to-failure traces is {model_summary['evidence']['nasa_late_exponent_raw']:.4f}, supporting an accelerated end-of-life regime.",
        f"The derived condition term is dominated by feed sensitivity ({model_summary['selected_coefficients']['feed_exponent']:.3f}), followed by depth ({model_summary['selected_coefficients']['depth_exponent']:.3f}); spindle-speed sensitivity is weaker ({model_summary['selected_coefficients']['speed_exponent']:.3f}) in the available data range.",
        "The selected final model is piecewise rather than single-regime: early wear is sub-linear in time, while late wear accelerates.",
        "A closed-form simulator was implemented so tool life can be predicted analytically for any target wear threshold once a calibration factor is chosen.",
    ]
    for item in summary_points:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("2. Data Base Used", level=1)
    document.add_paragraph(
        "Three public milling datasets available in the local workspace were used for complementary roles: "
        "NUAA to identify the process-parameter law, PHM 2010 to verify the early-regime exponent on a second milling source, "
        "and NASA to identify the accelerated late wear regime."
    )
    add_df_table(document, dataset_summary, decimals=2)

    document.add_heading("3. Selected Equation", level=1)
    p = document.add_paragraph()
    p.add_run("Condition intensity:\n").bold = True
    p.add_run(
        "phi = (speed_rpm / 1800)^0.3600 * (feed_mm_tooth / 0.05)^3.9560 * (depth_mm / 3.0)^0.6881"
    )
    p = document.add_paragraph()
    p.add_run("Early regime (VB <= 0.25 mm):\n").bold = True
    p.add_run("VB(t) = VB0 + lambda * 0.01858 * phi * t^0.64")
    p = document.add_paragraph()
    p.add_run("Late regime (VB > 0.25 mm):\n").bold = True
    p.add_run("VB(t) = 0.25 + lambda * rho * 0.01858 * phi * (t - t0.25)^1.29")

    document.add_paragraph(
        "This structure was selected because forcing a single exponent over the full milling life cycle underfits the early stage "
        "and misrepresents end-of-life acceleration. The transition wear of 0.25 mm is treated as the regime boundary."
    )
    add_df_table(document, key_results, decimals=4)

    document.add_heading("4. Fitted Wear Curves", level=1)
    document.add_paragraph(
        "The figures below show how the selected power-law family tracks the measured wear trajectories in each dataset role."
    )
    add_figure(document, PLOTS / "nuaa_early_regime_fit.png", "Figure 1. NUAA early-regime fit used for the condition law.")
    add_figure(document, PLOTS / "phm_early_regime_fit.png", "Figure 2. PHM 2010 fit used to confirm early-regime behavior.")
    add_figure(document, PLOTS / "nasa_late_regime_fit.png", "Figure 3. NASA fit used to identify late-stage acceleration.")

    document.add_heading("5. Deeper Validation", level=1)
    document.add_paragraph(
        "A deeper validation layer was added beyond the original phase-2 fit: bootstrap stability of the exponents, "
        "leave-one-experiment-out benchmarking against machine-learning baselines, and material-family analysis of late-stage acceleration."
    )
    add_figure(document, PLOTS / "bootstrap_exponent_stability.png", "Figure 4. Bootstrap distribution of early and late fitted exponents.")
    add_figure(document, PLOTS / "nuaa_loeo_benchmark.png", "Figure 5. NUAA leave-one-experiment-out benchmark across interpretable and ML models.")
    add_figure(document, PLOTS / "nuaa_parity_equation_vs_linear.png", "Figure 6. Parity comparison of the explicit equation and the best ML baseline.")
    add_df_table(document, benchmark_metrics, decimals=4)
    document.add_paragraph(
        "Interpretation: the explicit equation is not the numerically best pointwise predictor on held-out NUAA trajectories; "
        "a plain linear regression achieves lower error in this benchmark. However, the explicit equation remains the preferred phase-2 model "
        "because it is analytically invertible, exposes the machining parameters directly, and can be used as a simulation engine rather than only as a black-box regressor."
    )
    add_figure(document, PLOTS / "nasa_late_stage_materials.png", "Figure 7. Late-stage acceleration ratio grouped by material family in the NASA dataset.")
    add_df_table(document, late_ratio_summary, decimals=4)

    document.add_heading("6. Simulator Behavior", level=1)
    document.add_paragraph(
        "The simulator was built on top of the selected piecewise equation. The baseline example below corresponds to "
        "speed = 1800 rpm, feed = 0.05 mm/tooth, depth = 3.0 mm, initial wear = 0.05 mm, threshold = 0.30 mm, and calibration factor = 1.0."
    )
    sim = MillingToolLifeModel.from_json()
    baseline_life = sim.life_to_threshold_minutes(
        speed_rpm=1800.0,
        feed_mm_tooth=0.05,
        depth_mm=3.0,
        threshold_wear_mm=0.30,
        initial_wear_mm=0.05,
        calibration_factor=1.0,
        material_family="generic",
    )
    cast_iron_life = sim.life_to_threshold_minutes(
        speed_rpm=1800.0,
        feed_mm_tooth=0.05,
        depth_mm=3.0,
        threshold_wear_mm=0.30,
        initial_wear_mm=0.05,
        calibration_factor=1.0,
        material_family="cast_iron",
    )
    steel_life = sim.life_to_threshold_minutes(
        speed_rpm=1800.0,
        feed_mm_tooth=0.05,
        depth_mm=3.0,
        threshold_wear_mm=0.30,
        initial_wear_mm=0.05,
        calibration_factor=1.0,
        material_family="steel",
    )
    document.add_paragraph(f"Predicted generic baseline life: {baseline_life:.2f} min", style="List Bullet")
    document.add_paragraph(f"Predicted cast-iron baseline life: {cast_iron_life:.2f} min", style="List Bullet")
    document.add_paragraph(f"Predicted steel baseline life: {steel_life:.2f} min", style="List Bullet")
    add_figure(document, PLOTS / "simulation_example_curves.png", "Figure 8. Example simulated wear curves under low, medium, and high load.")
    add_figure(document, PLOTS / "life_sensitivity_curves.png", "Figure 9. Predicted tool-life sensitivity around the baseline condition.")

    document.add_heading("7. Findings", level=1)
    findings = [
        "The most stable observation across datasets is the existence of two wear regimes rather than one universal Taylor-like regime.",
        "Feed per tooth is the strongest driver in the fitted early condition law. Within this dataset family, increasing feed rapidly shortens life.",
        "Depth of cut has a clear but secondary influence. Spindle speed is less identifiable because the NUAA speed range is relatively narrow.",
        "Late-stage acceleration differs strongly by material family in the NASA cases, so a single absolute post-threshold constant is not justified.",
        "The explicit equation is best suited as a hybrid engineering model: interpretable, calibratable, and analytically invertible."
    ]
    for item in findings:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("8. Limits and Recommended Next Work", level=1)
    next_steps = [
        "Collect a wider spindle-speed sweep under otherwise fixed conditions. This is the cleanest way to stabilize the speed exponent.",
        "Add true run-to-failure data at the same parameter grid used for the early-regime fit. That would remove the current split between NUAA/PHM early evidence and NASA late evidence.",
        "Estimate the calibration factor lambda separately for each tool-workpiece pair before using the simulator in production planning.",
        "If maximum predictive accuracy is the only objective, keep a parallel ML baseline. If interpretability and simulation matter, keep the explicit equation as the main model."
    ]
    for item in next_steps:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("9. References", level=1)
    refs = [
        "PHM Society. 2010 PHM Society Conference Data Challenge. https://phmsociety.org/phm_competition/2010-phm-society-conference-data-challenge/",
        "NASA Ames / UC Berkeley. Documentation for Mill Data Set. Local copy used from data/nasa_milling/Readme.pdf.",
        "Piecuch, G., & Żabiński, T. (2025). A new open dataset from a milling process – data for classification and estimation of tool life. Scientific Data, 12, 650. https://www.nature.com/articles/s41597-025-04923-y",
        "Paszkiewicz, A. et al. (2023). Estimation of Tool Life in the Milling Process—Testing Regression Models. Sensors, 23(23), 9346. https://doi.org/10.3390/s23239346",
    ]
    for ref in refs:
        document.add_paragraph(ref, style="List Bullet")

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Appendix: Output Inventory", level=1)
    document.add_paragraph(
        "Primary generated artifacts are stored in the phase-2 output directory: coefficients JSON, CSV tables, plots, and this Word report."
    )

    document.save(REPORT_PATH)


def main() -> None:
    research_pipeline.main()

    nuaa_runs = load_nuaa_run_level()
    phm = load_phm_cut_level()
    nasa = load_nasa_case_level()

    dataset_profiles = pd.read_csv(OUTPUTS / "dataset_profiles.csv")
    dataset_summary = aggregate_dataset_summary(dataset_profiles)
    key_results = pd.read_csv(OUTPUTS / "key_results.csv")
    nasa_ratio_summary = pd.read_csv(OUTPUTS / "nasa_late_stage_ratio_summary.csv")
    nasa_ratio_df = pd.read_csv(OUTPUTS / "nasa_late_stage_case_ratios.csv")
    model_summary = json.loads((OUTPUTS / "phase2_model_summary.json").read_text(encoding="utf-8"))

    early_nuaa_boot = bootstrap_exponents(
        df=nuaa_runs,
        group_col="experiment_tag",
        time_col="relative_time_min",
        wear_col="wear_mm",
        n_boot=250,
        seed=42,
        initial_guess_log_a=-2.0,
        initial_guess_m=0.7,
    )
    early_phm_boot = bootstrap_exponents(
        df=phm,
        group_col="experiment_tag",
        time_col="relative_cut",
        wear_col="wear_mm",
        n_boot=250,
        seed=43,
        initial_guess_log_a=-5.0,
        initial_guess_m=0.8,
    )
    late_nasa_boot = bootstrap_exponents(
        df=nasa.rename(columns={"case": "case_id", "time": "time_min", "VB": "wear_mm"}),
        group_col="case_id",
        time_col="relative_time_min",
        wear_col="wear_mm",
        n_boot=250,
        seed=44,
        initial_guess_log_a=-3.0,
        initial_guess_m=1.1,
    )

    early_nuaa_boot.to_csv(OUTPUTS / "bootstrap_nuaa_early_exponent.csv", index=False)
    early_phm_boot.to_csv(OUTPUTS / "bootstrap_phm_early_exponent.csv", index=False)
    late_nasa_boot.to_csv(OUTPUTS / "bootstrap_nasa_late_exponent.csv", index=False)
    plot_bootstrap_results(
        early_nuaa=early_nuaa_boot,
        early_phm=early_phm_boot,
        late_nasa=late_nasa_boot,
        output_path=PLOTS / "bootstrap_exponent_stability.png",
    )

    benchmark_metrics, benchmark_predictions = benchmark_nuaa_models(nuaa_runs=nuaa_runs)
    benchmark_metrics.to_csv(OUTPUTS / "nuaa_loeo_benchmark_metrics.csv", index=False)
    benchmark_predictions.to_csv(OUTPUTS / "nuaa_loeo_benchmark_predictions.csv", index=False)
    plot_benchmark_results(benchmark_metrics, PLOTS / "nuaa_loeo_benchmark.png")
    plot_parity_results(benchmark_predictions, PLOTS / "nuaa_parity_equation_vs_linear.png")

    simulator = MillingToolLifeModel.from_json()
    plot_sensitivity_curves(simulator, PLOTS / "life_sensitivity_curves.png")
    plot_late_stage_materials(nasa_ratio_df, PLOTS / "nasa_late_stage_materials.png")

    build_report(
        model_summary=model_summary,
        dataset_summary=dataset_summary,
        key_results=key_results,
        benchmark_metrics=benchmark_metrics,
        late_ratio_summary=nasa_ratio_summary,
    )


if __name__ == "__main__":
    main()
